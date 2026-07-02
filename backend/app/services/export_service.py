"""
Export Service — compiles GTM launch packages into PDF, Word (.docx), Markdown, HTML, and CSV.
Includes Certificate generator and sub-component builders.
"""
import io
import csv
from typing import Any
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas


class ExportService:
    @staticmethod
    def generate_markdown(package: dict[str, Any]) -> str:
        """Compiles package outputs into formatted Markdown."""
        theme = package.get("theme", "N/A")
        domain = package.get("domain", "N/A")
        score = package.get("launch_readiness_score", 0)
        gtm = package.get("gtm_package", {})

        brand = gtm.get("brand", {})
        structure = gtm.get("structure", {})
        content = gtm.get("content", {})
        marketing = gtm.get("marketing", {})
        emails = gtm.get("email", {})
        execution = gtm.get("execution", {})

        md = []
        md.append(f"# Hackathon Launch Brief: {brand.get('event_name', 'Untitled Event')}")
        md.append(f"**Theme:** {theme} | **Domain:** {domain} | **Launch Score:** {score}%")
        md.append("\n" + "─" * 40 + "\n")

        # 1. Brand Guide
        md.append("## 🎨 1. Brand Identity Guide")
        md.append(f"- **Tagline:** {brand.get('tagline', 'N/A')}")
        md.append(f"- **Primary Color:** `{brand.get('color_primary', '#3d5eff')}`")
        md.append(f"- **Secondary Color:** `{brand.get('color_secondary', '#d946ef')}`")
        md.append(f"- **Tone Guideline:** {', '.join(brand.get('tone_adjectives', ['Professional']))}")
        md.append("\n**Target Audience Persona:**")
        md.append(brand.get("persona_text", "N/A"))
        md.append("\n")

        # 2. Structure & schedule
        md.append("## 📅 2. Schedule & Judging Criteria")
        md.append("### Hour-by-Hour Schedule")
        for item in structure.get("schedule", []):
            md.append(f"- **{item.get('time', 'N/A')}**: {item.get('activity', 'N/A')} ({item.get('duration_mins', 0)} mins) [{item.get('type', 'session')}]")
        
        md.append("\n### Judging Parameters")
        for crit in structure.get("judging_criteria", []):
            md.append(f"- **{crit.get('name', 'N/A')} ({crit.get('weight', 0)}%)**: {crit.get('description', '')}")
        md.append("\n")

        # 3. Copy & FAQ
        md.append("## 📝 3. Landing Page Copy & FAQ")
        md.append(f"### Hero Headline: **{content.get('hero_headline', '')}**")
        md.append(f"### Subheadline: *{content.get('hero_subheadline', '')}*")
        md.append("\n**About Section text:**")
        md.append(content.get("about_copy", ""))
        md.append("\n### FAQs")
        for faq in content.get("faq", []):
            md.append(f"**Q: {faq.get('question')}**")
            md.append(f"A: {faq.get('answer')}\n")

        # 4. Social marketing
        md.append("## 📢 4. Social Marketing Package")
        md.append("### LinkedIn Post")
        md.append(marketing.get("linkedin_announcement", "N/A"))
        md.append("\n### Twitter Announcement Thread")
        for tweet in marketing.get("twitter_thread", []):
            md.append(f"Tweet {tweet.get('tweet_num', 1)}: {tweet.get('content')}")
        md.append("\n")

        # 5. Emails
        md.append("## ✉️ 5. Email Sequences")
        md.append("### Builder invitation")
        md.append(f"**Subject:** {emails.get('invite_subject')}\n")
        md.append(emails.get("invite_body", ""))
        md.append("\n### Sponsor outreach")
        md.append(f"**Subject:** {emails.get('sponsor_subject')}\n")
        md.append(emails.get("sponsor_body", ""))
        md.append("\n")

        # 6. Operations & budget
        md.append("## ⚙️ 6. Operations & Risk Roadmap")
        for week in execution.get("weekly_plan", []):
            md.append(f"### Week {week.get('week')}")
            for t in week.get("tasks", []):
                md.append(f"- [ ] {t}")
        md.append("\n### Mitigation Strategy")
        for r in execution.get("risk_plan", []):
            md.append(f"- **Risk:** {r.get('risk')}\n  **Mitigation:** {r.get('mitigation')}")

        return "\n".join(md)

    @staticmethod
    def generate_html(package: dict[str, Any]) -> str:
        """Formats package details inside styled HTML pages."""
        md = ExportService.generate_markdown(package)
        # Parse markdown headers/lists to basic HTML
        html_body = md.replace("\n", "<br>").replace("# ", "<h1>").replace("## ", "<h2>").replace("### ", "<h3>")
        
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Hackathon GTM Package</title>
            <style>
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
                    line-height: 1.6;
                    color: #1f2937;
                    max-width: 800px;
                    margin: 40px auto;
                    padding: 0 20px;
                    background-color: #f9fafb;
                }}
                h1 {{ color: #1e3a8a; border-bottom: 2px solid #e5e7eb; padding-bottom: 10px; }}
                h2 {{ color: #2563eb; margin-top: 30px; }}
                h3 {{ color: #1d4ed8; }}
                pre, code {{ background-color: #f3f4f6; padding: 2px 6px; border-radius: 4px; font-family: monospace; }}
                br {{ margin-bottom: 10px; }}
            </style>
        </head>
        <body>
            {html_body}
        </body>
        </html>
        """

    @staticmethod
    def generate_csv(package: dict[str, Any]) -> str:
        """Renders schedule activities and budget rows into a CSV string."""
        gtm = package.get("gtm_package", {})
        structure = gtm.get("structure", {})
        schedule = structure.get("schedule", [])

        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["Time Slot", "Activity Name", "Duration (Mins)", "Category Type"])

        for item in schedule:
            writer.writerow([
                item.get("time", ""),
                item.get("activity", ""),
                item.get("duration_mins", 0),
                item.get("type", "session")
            ])

        return output.getvalue()

    @staticmethod
    def generate_docx(package: dict[str, Any]) -> io.BytesIO:
        """Generates Word (.docx) package containing styled text and tables."""
        theme = package.get("theme", "N/A")
        domain = package.get("domain", "N/A")
        gtm = package.get("gtm_package", {})

        brand = gtm.get("brand", {})
        structure = gtm.get("structure", {})
        content = gtm.get("content", {})
        marketing = gtm.get("marketing", {})
        emails = gtm.get("email", {})
        execution = gtm.get("execution", {})

        doc = Document()
        doc.add_heading(brand.get("event_name", "Hackathon GTM Package"), 0)

        p = doc.add_paragraph()
        p.add_run(f"Theme: {theme} | Domain: {domain}\n").bold = True

        doc.add_heading("1. Brand & Tagline", level=1)
        doc.add_paragraph(f"Tagline: {brand.get('tagline', 'N/A')}")
        doc.add_paragraph(brand.get("persona_text", ""))

        doc.add_heading("2. Schedule & Timeline", level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Time"
        hdr_cells[1].text = "Activity"
        hdr_cells[2].text = "Duration"
        hdr_cells[3].text = "Category"

        for item in structure.get("schedule", []):
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get("time", ""))
            row_cells[1].text = str(item.get("activity", ""))
            row_cells[2].text = f"{item.get('duration_mins', 0)} mins"
            row_cells[3].text = str(item.get("type", ""))

        doc.add_heading("3. Landing Page Copy", level=1)
        doc.add_paragraph(content.get("hero_headline", ""))
        doc.add_paragraph(content.get("about_copy", ""))

        doc.add_heading("4. Marketing social copy", level=1)
        doc.add_paragraph("LinkedIn:")
        doc.add_paragraph(marketing.get("linkedin_announcement", ""))

        # Save to bytes stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    @staticmethod
    def generate_pdf(package: dict[str, Any]) -> io.BytesIO:
        """Generates PDF booklet package using ReportLab."""
        gtm = package.get("gtm_package", {})
        brand = gtm.get("brand", {})
        structure = gtm.get("structure", {})
        content = gtm.get("content", {})

        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Title"],
            fontSize=24,
            textColor=colors.HexColor("#1e3a8a"),
            spaceAfter=20
        )
        h1_style = ParagraphStyle(
            "DocH1",
            parent=styles["Heading1"],
            fontSize=16,
            textColor=colors.HexColor("#2563eb"),
            spaceBefore=15,
            spaceAfter=10
        )
        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["BodyText"],
            fontSize=10,
            leading=14,
            spaceAfter=8
        )

        story = []
        
        # Cover / Header
        story.append(Paragraph(brand.get("event_name", "Hackathon GTM Package"), title_style))
        story.append(Paragraph(f"Tagline: {brand.get('tagline', '')}", body_style))
        story.append(Spacer(1, 15))

        # Brand
        story.append(Paragraph("1. Brand Guide", h1_style))
        story.append(Paragraph(brand.get("persona_text", ""), body_style))
        story.append(Spacer(1, 10))

        # Landing Copy
        story.append(Paragraph("2. Landing Page Copy", h1_style))
        story.append(Paragraph(f"Headline: {content.get('hero_headline', '')}", body_style))
        story.append(Paragraph(content.get("about_copy", ""), body_style))
        story.append(Spacer(1, 10))

        # Schedule Table
        story.append(Paragraph("3. Hourly Timetable", h1_style))
        table_data = [["Time Slot", "Activity Name", "Duration", "Type"]]
        for item in structure.get("schedule", []):
            table_data.append([
                str(item.get("time")),
                str(item.get("activity")),
                f"{item.get('duration_mins')} mins",
                str(item.get("type"))
            ])
        
        t = Table(table_data, colWidths=[80, 200, 80, 80])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#f3f4f6")),
            ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor("#1f2937")),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 9),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e5e7eb")),
            ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
            ('FONTSIZE', (0,1), (-1,-1), 8),
        ]))
        story.append(t)

        doc.build(story)
        pdf_buffer.seek(0)
        return pdf_buffer

    @staticmethod
    def generate_certificate(event_name: str, recipient_name: str = "Hackathon Competitor", award_title: str = "Certificate of Participation") -> io.BytesIO:
        """Generates participant/winner award certificate using ReportLab canvas."""
        pdf_buffer = io.BytesIO()
        # Certificate page in Landscape letter format (11 x 8.5 inches = 792 x 612 points)
        c = canvas.Canvas(pdf_buffer, pagesize=(792, 612))

        # Draw decorative border lines
        c.setStrokeColor(colors.HexColor("#1e3a8a"))
        c.setLineWidth(4)
        c.rect(30, 30, 732, 552)

        c.setStrokeColor(colors.HexColor("#d946ef"))
        c.setLineWidth(1.5)
        c.rect(38, 38, 716, 536)

        # Draw branding text
        c.setFont("Helvetica-Bold", 32)
        c.setFillColor(colors.HexColor("#1e3a8a"))
        c.drawCentredString(396, 450, "CERTIFICATE OF HONOR")

        c.setFont("Helvetica", 14)
        c.setFillColor(colors.HexColor("#4b5563"))
        c.drawCentredString(396, 400, "This is proudly presented to")

        # Recipient name
        c.setFont("Helvetica-Bold", 26)
        c.setFillColor(colors.HexColor("#0f172a"))
        c.drawCentredString(396, 340, recipient_name)

        # Award Title
        c.setFont("Helvetica-Oblique", 14)
        c.setFillColor(colors.HexColor("#4b5563"))
        c.drawCentredString(396, 290, f"for completing the GTM campaign roadmap of")

        # Event Name
        c.setFont("Helvetica-Bold", 20)
        c.setFillColor(colors.HexColor("#2563eb"))
        c.drawCentredString(396, 240, event_name)

        # Signature
        c.setFont("Helvetica-Bold", 10)
        c.setFillColor(colors.HexColor("#1f2937"))
        c.drawString(100, 120, "HACKLAUNCH AI SIGNATURE")
        c.setStrokeColor(colors.HexColor("#9ca3af"))
        c.setLineWidth(1)
        c.line(100, 140, 280, 140)

        c.drawCentredString(600, 120, "DATE OF ISSUE")
        c.line(520, 140, 680, 140)
        c.drawString(560, 145, "June 30, 2026")

        c.showPage()
        c.save()

        pdf_buffer.seek(0)
        return pdf_buffer
